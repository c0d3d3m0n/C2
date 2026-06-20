"""Generate a detailed DOCX project report for this repository."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parent
DEFAULT_OUTPUT = PROJECT_ROOT / "project_report.docx"


def collect_project_tree(root: Path) -> list[str]:
    items: list[str] = []
    for path in sorted(root.rglob("*")):
        if path.is_dir():
            continue

        relative = path.relative_to(root)
        parts = relative.parts
        if not parts:
            continue
        if parts[0] in {"venv", ".git", "__pycache__"}:
            continue
        if path.suffix in {".pyc", ".db"}:
            continue
        if len(parts) == 1 or parts[0] in {"agent", "client", "server"}:
            items.append(str(relative).replace("\\", "/"))
    return items


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.add_run(text)


def add_section(document: Document, title: str, paragraphs: list[str], bullets: list[str] | None = None) -> None:
    document.add_heading(title, level=1)
    for paragraph_text in paragraphs:
        document.add_paragraph(paragraph_text)
    for bullet in bullets or []:
        add_bullet(document, bullet)


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(6)


def build_document(output_path: Path) -> None:
    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("C2 Framework Project Report")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    document.add_paragraph(
        "This report summarizes the repository structure, runtime design, API surface, data model, deployment notes, and usage patterns."
    )

    add_section(
        document,
        "Executive Summary",
        [
            "The repository contains a compact FastAPI-based C2-style framework split into server, client, and agent components.",
            "The server stores agents, tasks, and results in SQLite, while the client exposes operator commands and the agent performs beaconing and task execution.",
        ],
    )

    add_section(
        document,
        "Repository Structure",
        ["The project is organized around a minimal three-part architecture with a few supporting documentation and deployment files."],
        [
            "server/ - FastAPI server, models, schemas, and database configuration.",
            "client/ - Operator CLI for queueing tasks and inspecting agent state.",
            "agent/ - Polling implant that registers, receives tasks, and submits results.",
            "requirements.txt - Python dependencies used by the project.",
            "Dockerfile and render.yaml - Container and deployment configuration.",
        ],
    )

    add_section(
        document,
        "Runtime Flow",
        [
            "1. The agent registers with the server and receives a UUID identifier.",
            "2. The agent periodically beacons to ask for queued tasks.",
            "3. The server returns pending work and marks tasks as sent.",
            "4. The agent executes the task locally and posts the result back.",
            "5. The client CLI queries agents, lists tasks, and retrieves results.",
        ],
    )

    add_section(
        document,
        "Server API",
        ["The FastAPI application exposes endpoints for registration, beaconing, task creation, result submission, and operator listing queries."],
        [
            "GET / - health check.",
            "POST /api/register - create an agent record.",
            "POST /api/beacon/{agent_id} - fetch queued tasks and update last-seen time.",
            "POST /api/results - store task output and mark completion.",
            "GET /api/agents - list all agents.",
            "POST /api/tasks/{agent_id} - queue a task for an agent.",
            "GET /api/tasks/{agent_id} - list tasks for an agent.",
            "GET /api/results/{task_id} - fetch the stored result for a task.",
        ],
    )

    add_section(
        document,
        "Database Model",
        ["The project uses SQLAlchemy ORM models with an SQLite backend stored in c2.db."],
        [
            "Agent: hostname, ip, os, last_seen, status, and generated UUID primary key.",
            "Task: command, arguments, status, creation timestamp, and agent foreign key.",
            "Result: task output and execution timestamp linked to a task.",
        ],
    )

    add_section(
        document,
        "Operator CLI",
        ["The client CLI wraps the REST API and provides a simple workflow for managing agents and queued tasks."],
        [
            "agents - list active agents.",
            "tasks <agent_id> - list queued and completed tasks.",
            "exec <agent_id> <command> [args] - queue a shell execution task.",
            "upload <agent_id> <local_file> <remote_path> - queue a file write task.",
            "download <agent_id> <remote_path> <local_file> - queue a file read task.",
            "get-file <task_id> <local_file> - save a completed download result locally.",
        ],
    )

    add_section(
        document,
        "Agent Behavior",
        ["The agent implements a beacon loop with jitter and supports a small task vocabulary including shell execution, file transfer, and sleep adjustment."],
        [
            "Registers using host metadata gathered from socket and platform modules.",
            "Polls the beacon endpoint on a configurable interval.",
            "Executes tasks and posts results back to the server.",
            "Supports base64-encoded file transfer payloads.",
        ],
    )

    add_section(
        document,
        "Installation And Local Run",
        ["A typical local setup uses a Python virtual environment, installs dependencies, starts the FastAPI server, and then runs the agent and CLI in separate terminals."],
    )
    add_code_block(
        document,
        "python -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt\nuvicorn server.main:app --reload\npython agent/implant.py\npython client/cli.py agents",
    )

    add_section(
        document,
        "Deployment Notes",
        [
            "The repository includes Docker and Render configuration for containerized deployment.",
            "SQLite is convenient for local use but is not durable on ephemeral infrastructure without an external database or persistent disk.",
        ],
    )

    add_section(
        document,
        "Project Tree Snapshot",
        ["The following paths were collected automatically from the repository:"],
        collect_project_tree(PROJECT_ROOT),
    )

    add_section(
        document,
        "Recommended Improvements",
        ["The codebase would benefit from authentication, stronger persistence, structured logging, and test coverage for the API and CLI paths."],
        [
            "Replace SQLite with a production database for durable deployments.",
            "Add API authentication and authorization.",
            "Add tests for endpoints, task execution, and CLI flows.",
            "Split transport, persistence, and task execution into smaller modules as the project grows.",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a DOCX project report for this repository.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="Output DOCX file path")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_path = Path(args.output).resolve()
    build_document(output_path)
    print(f"Project report written to {output_path}")


if __name__ == "__main__":
    main()